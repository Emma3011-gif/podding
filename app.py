#!/usr/bin/env python
"""
Unified Flask Application - Combines backend API and frontend serving
Run with: python app.py
"""

from flask import Flask, render_template, request, jsonify, session, Response, stream_with_context, redirect, url_for
import requests
import os
import uuid
import time
import numpy as np
import io
from PyPDF2 import PdfReader
from openai import OpenAI
from flask_cors import CORS
from dotenv import load_dotenv
from pathlib import Path
from PIL import Image
import pytesseract
from docx import Document
import tempfile

# Load environment variables IMMEDIATELY before any local module imports
env_path = Path(__file__).parent / '.env'
load_dotenv(dotenv_path=env_path)

# Helper function to make strings safe for Windows console output (cp1252 encoding)
def safe_str(s, max_len=100):
    """Convert string to ASCII-safe version for console output, stripping non-ASCII characters"""
    if not s:
        return ""
    try:
        # Try to encode as cp1252 (Windows default) - if it fails, fall back to ASCII
        s.encode('cp1252')
        return s[:max_len]
    except (UnicodeEncodeError, UnicodeDecodeError):
        # Fallback: encode to ASCII, ignoring all non-ASCII characters (including emojis)
        return s.encode('ascii', errors='ignore').decode('ascii')[:max_len]

# Configuration
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "").strip()
OPENROUTER_MODEL = os.getenv("OPENROUTER_MODEL", "openai/gpt-4o-mini")
ENV_MODE = os.getenv("ENV", "production")
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", 1000))
TOP_K_RESULTS = int(os.getenv("TOP_K_RESULTS", 5))  # Increased from 3 to 5 for better context
DEBUG_AUTH = os.getenv('DEBUG_AUTH', 'false').lower() == 'true'  # Debug flag for auth endpoints

# Session security - Fixed secret key for production (required for session persistence)
SECRET_KEY = os.getenv('SECRET_KEY')
if not SECRET_KEY:
    if os.getenv('VERCEL') or ENV_MODE == 'production':
        print("[ERROR] SECRET_KEY environment variable is required in production!")
        # Generate a warning but allow fallback for compatibility
        SECRET_KEY = os.urandom(24)
        print("[WARN] Using temporary secret key - sessions will not persist across cold starts!")
    else:
        # Local development: use a fixed default for convenience
        SECRET_KEY = 'dev-secret-key-change-in-production-12345'
        print("[INFO] Using development SECRET_KEY (set SECRET_KEY env var for production)")

# Import database models AFTER environment is loaded
try:
    import models  # Our new database models
    MODELS_AVAILABLE = True
except ImportError as e:
    print(f"[WARN] Failed to import models: {e}")
    MODELS_AVAILABLE = False

# Detect if Tesseract OCR is available
TESSERACT_AVAILABLE = True
if os.getenv('VERCEL') or os.getenv('TESSERACT_DISABLED'):
    TESSERACT_AVAILABLE = False
    print("[CONFIG] Tesseract OCR disabled (Vercel/serverless environment or explicitly disabled)")
else:
    try:
        pytesseract.get_tesseract_version()
        print("[OK] Tesseract OCR available")
    except Exception:
        TESSERACT_AVAILABLE = False
        print("[WARN] Tesseract OCR not available - image text extraction will be disabled")

# Configure pytesseract path if TESSERACT_CMD is set
tesseract_cmd = os.getenv("Tesseract_CMD") or os.getenv("TESSERACT_CMD")
if tesseract_cmd:
    pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
    print(f"[CONFIG] Tesseract OCR configured: {tesseract_cmd}")
else:
    print("[INFO] Tesseract OCR: using system PATH (ensure tesseract is installed)")

# Test Tesseract availability at startup
def check_tesseract():
    """Verify Tesseract is installed and accessible"""
    try:
        # Try to get tesseract version
        version = pytesseract.get_tesseract_version()
        print(f"[OK] Tesseract OCR ready (version {version})")
        return True
    except Exception as e:
        print(f"[ERROR] Tesseract OCR check failed: {e}")
        print("[WARN] Image uploads will fail. Install Tesseract or set TESSERACT_CMD in .env")
        return False

# Initialize Flask app
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Determine template and static folders with comprehensive fallbacks
template_folder = None
static_folder = None

print("[CONFIG] === Template/Static folder detection ===")
print(f"[CONFIG] BASE_DIR: {BASE_DIR}")
print(f"[CONFIG] CWD: {os.getcwd()}")

# Build comprehensive list of locations to check
locations_to_check = []

# 1. Check AWS Lambda / serverless typical layouts
if '/var/task' in BASE_DIR or os.getcwd().startswith('/var/task'):
    locations_to_check.extend([
        '/var/task/templates',
        '/var/task/app/templates',
        '/var/task/./templates',
        os.path.join(os.getcwd(), 'templates'),
    ])

# 2. Check relative to app.py
locations_to_check.extend([
    os.path.join(BASE_DIR, 'templates'),
    os.path.join(BASE_DIR, '..', 'templates'),
    os.path.join(BASE_DIR, '..', '..', 'templates'),
    os.path.join(BASE_DIR, '..', '..', '..', 'templates'),
])

# 3. Check relative to current working directory
locations_to_check.extend([
    os.path.join(os.getcwd(), 'templates'),
    'templates',
])

# 4. Check absolute common paths
locations_to_check.extend([
    '/app/templates',
    '/workspace/templates',
])

# Remove duplicates while preserving order
seen = set()
template_locations = []
for loc in locations_to_check:
    abs_loc = os.path.abspath(loc)
    if abs_loc not in seen:
        seen.add(abs_loc)
        template_locations.append(loc)

print(f"[CONFIG] Checking {len(template_locations)} template locations...")

# Find first existing template folder
for loc in template_locations:
    if os.path.exists(loc) and os.path.isdir(loc):
        # Check if it has at least one .html file
        try:
            files = os.listdir(loc)
            html_files = [f for f in files if f.endswith('.html')]
            if html_files:
                template_folder = loc
                print(f"[CONFIG] [OK] Found templates at: {os.path.abspath(loc)}")
                print(f"[CONFIG]   HTML files: {html_files[:5]}{'...' if len(html_files) > 5 else ''}")
                break
            else:
                print(f"[CONFIG] [SKIP] '{loc}' - no .html files found")
        except Exception as e:
            print(f"[CONFIG] [ERROR] Error reading '{loc}': {e}")

# Do the same for static folder
static_locations = []
for loc in locations_to_check:
    static_loc = loc.replace('templates', 'static')
    if static_loc not in [l.replace('templates', 'static') for l in static_locations]:
        static_locations.append(static_loc)

for loc in static_locations:
    if os.path.exists(loc) and os.path.isdir(loc):
        static_folder = loc
        print(f"[CONFIG] [OK] Found static at: {os.path.abspath(loc)}")
        break

# If still not found, try to locate templates by searching nearby directories
if not template_folder:
    print("[CONFIG] Primary search failed, trying fallback search...")
    # Look for any 'templates' directory within a reasonable range
    search_root = os.path.dirname(BASE_DIR)
    for root, dirs, files in os.walk(search_root, topdown=True, followlinks=False):
        if 'templates' in dirs:
            candidate = os.path.join(root, 'templates')
            try:
                files_in_candidate = os.listdir(candidate)
                html_files = [f for f in files_in_candidate if f.endswith('.html')]
                if html_files:
                    template_folder = candidate
                    print(f"[CONFIG] [OK] Found templates via walk: {os.path.abspath(candidate)}")
                    print(f"[CONFIG]   HTML files: {html_files[:5]}{'...' if len(html_files) > 5 else ''}")
                    break
            except Exception:
                continue
        # Limit depth to avoid excessive searching
        if root.count(os.sep) - search_root.count(os.sep) >= 3:
            dirs[:] = [d for d in dirs if d not in ['templates', 'static', 'venv', '__pycache__']]

# Final fallback - use defaults
if not template_folder:
    template_folder = os.path.join(BASE_DIR, 'templates')
    print(f"[WARN] Template folder NOT FOUND - defaulting to: {os.path.abspath(template_folder)}")
    print(f"[WARN]   This path may not exist or may not contain templates!")
else:
    # Verify critical templates exist
    auth_html = os.path.join(template_folder, 'auth.html')
    index_html = os.path.join(template_folder, 'index.html')
    if not os.path.exists(auth_html):
        print(f"[WARN]   auth.html NOT FOUND at {auth_html}")
    if not os.path.exists(index_html):
        print(f"[WARN]   index.html NOT FOUND at {index_html}")

if not static_folder:
    static_folder = os.path.join(BASE_DIR, 'static')
    print(f"[WARN] Static folder not found - defaulting to: {os.path.abspath(static_folder)}")

print(f"[CONFIG] Final template_folder: {os.path.abspath(template_folder)}")
print(f"[CONFIG] Final static_folder: {os.path.abspath(static_folder)}")
print("[CONFIG] === End detection ===")

print(f"[CONFIG] BASE_DIR: {BASE_DIR}")
print(f"[CONFIG] CWD: {os.getcwd()}")
print(f"[CONFIG] template_folder: {template_folder}")
print(f"[CONFIG] static_folder: {static_folder}")

# Verify template folder contents
if os.path.exists(template_folder):
    print(f"[CONFIG] Templates found: {os.listdir(template_folder)}")
else:
    print(f"[WARN] Template folder does not exist: {template_folder}")

app = Flask(__name__,
            template_folder=template_folder,
            static_folder=static_folder)
app.secret_key = SECRET_KEY
CORS(app)

# Debug mode based on environment
debug_mode = ENV_MODE == "production"

# Global variables
client = None
EMBED_CACHE = {}  # Cache for embedding vectors of individual text chunks
DB = {}  # In-memory cache for document embeddings (loaded from database on demand)

# Initialize database on startup (with error handling for read-only filesystems)
try:
    models.init_db()
except Exception as e:
    print(f"[WARN] Database initialization failed: {e}")
    print("[WARN] The app may not work in serverless environments without a writable database.")

# ==================== AUTHENTICATION ====================

# Import auth blueprint
try:
    from auth_integration import auth_bp, login_required, init_oauth
    app.register_blueprint(auth_bp)
    init_oauth(app)  # Initialize OAuth with the Flask app
    print("[OK] Authentication enabled")
except ImportError as e:
    print(f"[WARN] Auth module not found: {e}")
    # Define dummy decorator if auth not available
    def login_required(view_func):
        from functools import wraps
        @wraps(view_func)
        def wrapped_view(*args, **kwargs):
            return view_func(*args, **kwargs)
        return wrapped_view

# ==================== BACKEND LOGIC ====================

def initialize_openrouter():
    """Initialize OpenRouter client"""
    global client
    if not OPENROUTER_API_KEY:
        print("ERROR: OPENROUTER_API_KEY is not set")
        client = None
    else:
        try:
            client = OpenAI(
                api_key=OPENROUTER_API_KEY,
                base_url="https://openrouter.ai/api/v1",
                timeout=30.0
            )
            print(f"[OK] OpenRouter client initialized")
            print(f"  Model: {OPENROUTER_MODEL}")
            print(f"  Mode: {ENV_MODE}")
        except Exception as e:
            print(f"ERROR: Failed to initialize OpenRouter: {e}")
            client = None

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes"""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        texts = []
        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    texts.append(page_text)
            except Exception as e:
                print(f"Warning: Could not extract text from page {i}: {e}")
                continue
        return "\n\n".join(texts)
    except Exception as e:
        print(f"Failed to read PDF: {e}")
        raise

def extract_text_from_image(image_bytes: bytes) -> str:
    """Extract text from image using Tesseract OCR (optional)"""
    if not TESSERACT_AVAILABLE:
        return ""  # OCR disabled

    try:
        image = Image.open(io.BytesIO(image_bytes))
        # Convert to RGB if necessary (handles RGBA, grayscale, etc.)
        if image.mode != 'RGB':
            image = image.convert('RGB')
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as e:
        print(f"Failed to read image with OCR: {e}")
        return ""  # Return empty text instead of raising

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n\n".join(paragraphs)
    except Exception as e:
        print(f"Failed to read DOCX: {e}")
        raise


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Unified text extraction: auto-detect file type and extract accordingly"""
    filename_lower = filename.lower()
    if filename_lower.endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith('.docx'):
        return extract_text_from_docx(file_bytes)
    elif any(filename_lower.endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.webp', '.bmp']):
        return extract_text_from_image(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}")

def chunk(text, size=None):
    """Split text into chunks"""
    if size is None:
        size = CHUNK_SIZE

    if len(text) <= size:
        return [text]

    chunks = []
    start = 0
    while start < len(text):
        end = start + size
        if end < len(text):
            # Try to find sentence boundary
            for i in range(end, start + max(50, size // 2), -1):
                if i < len(text) and text[i] in '.!?\n':
                    end = i + 1
                    break
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end

    return chunks

def generate_document_analysis(text: str, filename: str) -> str:
    """Generate an intelligent analysis of the document including main topics, key concepts, and structure"""
    try:
        if client is None:
            return "Document uploaded successfully! I'm ready to answer your questions about it."

        # Take first several chunks for analysis (to stay within token limits)
        # We'll use a representative sample of the document
        preview_text = text[:4000] if len(text) > 4000 else text

        analysis_prompt = f"""Analyze this document and provide a comprehensive overview. The filename is: {filename}

Document content (first part):
---
{preview_text}
---

Provide an analysis with these sections:

1. **Document Type & Purpose**: What kind of document is this? (e.g., research paper, textbook chapter, lecture notes, report, etc.)

2. **Main Topic**: What is the central subject or theme?

3. **Key Concepts & Terms**: List 5-8 important concepts, theories, or terminology introduced

4. **Document Structure**: How is the document organized? (e.g., sections, headings, flow)

5. **What You'll Find Inside**: What topics are covered? Give the user a mental map of the content

6. **Suggested Starting Points**: Recommend 2-3 good questions the user could ask to begin exploring this document

Be concise but informative. Use a friendly, welcoming tone. Format with clear headings."""

        response = client.chat.completions.create(
            model=OPENROUTER_MODEL,
            messages=[
                {"role": "system", "content": "You are an expert document analyst. Your job is to quickly read document excerpts and provide users with a clear, structured overview that helps them understand what the document contains and how to navigate it. Be helpful and educational."},
                {"role": "user", "content": analysis_prompt}
            ],
            temperature=0.7,
            max_tokens=1000
        )

        analysis = response.choices[0].message.content
        return analysis

    except Exception as e:
        print(f"Error generating document analysis: {e}")
        # Fallback simple message
        return f"Document '{filename}' uploaded successfully! The document contains {len(text)} characters. I'm ready to answer your questions about its content."

def embed(text):
    """Generate embedding for text with caching"""
    if client is None:
        raise RuntimeError("OpenRouter client not initialized")

    # Check cache first
    text_hash = hash(text)
    if text_hash in EMBED_CACHE:
        return EMBED_CACHE[text_hash]

    try:
        start = time.time()
        response = client.embeddings.create(
            model="text-embedding-3-small",
            input=text
        )
        elapsed = time.time() - start
        embedding = response.data[0].embedding
        EMBED_CACHE[text_hash] = embedding

        if elapsed > 1.0:
            print(f"[EMBED] {elapsed:.2f}s for {len(text)} chars")
        return embedding
    except Exception as e:
        error_msg = str(e).lower()
        if "401" in str(e) or "unauthorized" in error_msg:
            raise RuntimeError("OpenRouter authentication failed. Check your API key.")
        elif "insufficient_quota" in error_msg or "402" in str(e):
            raise RuntimeError("OpenRouter quota exceeded. Add credits to your account.")
        elif "rate limit" in error_msg or "429" in str(e):
            raise RuntimeError("Rate limit exceeded. Please wait.")
        elif "model not found" in error_msg or "404" in str(e):
            raise RuntimeError("Embedding model not available. Your account may not have access.")
        else:
            raise RuntimeError(f"Failed to generate embeddings: {str(e)}")

def cosine(a, b):
    """Cosine similarity"""
    return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))

# ==================== EMBEDDING PERSISTENCE HELPERS ====================

def get_doc_embeddings(doc_id):
    """Get embeddings for a document, loading from database if not in memory

    Returns:
        List of tuples [(chunk_text, embedding_vector), ...]

    Raises:
        KeyError: if document not found in database
    """
    if doc_id in DB:
        return DB[doc_id]

    # Load from database
    embeddings = models.load_embeddings(doc_id)
    if embeddings is None:
        raise KeyError(f"Document {doc_id} not found in database")

    DB[doc_id] = embeddings
    return embeddings

def save_doc_embeddings(doc_id, embeddings):
    """Save embeddings to both memory and database"""
    DB[doc_id] = embeddings
    models.save_embeddings(doc_id, embeddings)

def delete_doc_embeddings(doc_id):
    """Remove embeddings from memory and database"""
    if doc_id in DB:
        del DB[doc_id]
    models.delete_embeddings(doc_id)

# ==================== FLASK ROUTES ====================

@app.route('/')
@login_required
def index():
    """Serve the main HTML page (protected)"""
    try:
        return render_template('index.html')
    except Exception as e:
        import os
        error_details = {
            'error': 'Template not found',
            'exception': str(e),
            'template_folder': app.template_folder,
            'template_exists': os.path.exists(os.path.join(app.template_folder, 'index.html')),
            'folder_contents': os.listdir(app.template_folder) if os.path.exists(app.template_folder) else 'FOLDER MISSING',
        }
        print(f"[ERROR] Index template render failed: {error_details}")
        return f"<h1>Error</h1><pre>{error_details}</pre>", 500

@app.route('/status')
def status():
    """Check backend status"""
    return jsonify({
        'backend': 'connected',
        'mode': ENV_MODE,
        'model': OPENROUTER_MODEL,
        'url': 'unified'
    })

@app.route('/debug/templates')
def debug_templates():
    """Debug endpoint to show template configuration (only when DEBUG_AUTH is true)"""
    if not DEBUG_AUTH:
        return jsonify({'error': 'Debug endpoint disabled'}), 403

    import os
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

    info = {
        'debug': True,
        'base_dir': BASE_DIR,
        'cwd': os.getcwd(),
        'template_folder': app.template_folder,
        'static_folder': app.static_folder,
        'template_folder_exists': os.path.exists(app.template_folder),
        'static_folder_exists': os.path.exists(app.static_folder),
        'templates': [],
        'static_files': []
    }

    # List template files
    if os.path.exists(app.template_folder):
        try:
            info['templates'] = os.listdir(app.template_folder)
        except Exception as e:
            info['templates_error'] = str(e)

    # List static files (first level only)
    if os.path.exists(app.static_folder):
        try:
            info['static_files'] = os.listdir(app.static_folder)
        except Exception as e:
            info['static_error'] = str(e)

    # Check specific template files
    info['auth_html_exists'] = os.path.exists(os.path.join(app.template_folder, 'auth.html'))
    info['index_html_exists'] = os.path.exists(os.path.join(app.template_folder, 'index.html'))

    return jsonify(info)

@app.route('/auth/check')
def auth_check():
    """Check if user is authenticated"""
    if 'user_id' in session:
        return jsonify({
            'authenticated': True,
            'user': {
                'id': session.get('user_id'),
                'email': session.get('user_email'),
                'name': session.get('user_name')
            }
        })
    return jsonify({'authenticated': False}), 401

@app.route('/upload', methods=['POST'])
@login_required
def upload_pdf():
    """Upload and process PDF"""
    try:
        if client is None:
            return jsonify({'error': 'OpenRouter client not initialized. Check OPENROUTER_API_KEY.'}), 503

        if 'file' not in request.files:
            return jsonify({'error': 'No file selected'}), 400

        file = request.files['file']

        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        # Determine allowed extensions based on OCR availability
        allowed_extensions = {'.pdf', '.docx'}
        if TESSERACT_AVAILABLE:
            allowed_extensions.update({'.jpg', '.jpeg', '.png', '.webp', '.bmp'})

        if not any(file.filename.lower().endswith(ext) for ext in allowed_extensions):
            if TESSERACT_AVAILABLE:
                return jsonify({'error': 'Only PDF, DOCX, and image files (JPG, PNG, etc.) are allowed'}), 400
            else:
                return jsonify({'error': 'Only PDF and DOCX files are allowed. Image OCR is not available on this server.'}), 400

        # Clear embed cache for new document
        EMBED_CACHE.clear()
        print("[CACHE] Embedding cache cleared")

        content = file.read()

        if len(content) > 25 * 1024 * 1024:
            return jsonify({'error': 'File size exceeds 25MB limit'}), 400

        # Extract text
        t0 = time.time()
        text = extract_text(content, file.filename)
        t1 = time.time()

        if not text or len(text.strip()) == 0:
            ext = Path(file.filename).suffix.lower()
            if ext in {'.jpg', '.jpeg', '.png', '.webp', '.bmp'}:
                if not TESSERACT_AVAILABLE:
                    return jsonify({'error': 'Image OCR is not available on this server. Please upload a PDF or DOCX file.'}), 400
                else:
                    return jsonify({'error': 'Could not extract text from image. The image may be scanned, low quality, or contain no readable text.'}), 400
            else:
                return jsonify({'error': 'Could not extract text from document. The file may be corrupted or contain no readable text.'}), 400

        print(f"[PDF] Text extraction: {t1-t0:.2f}s, {len(text)} chars")

        # Chunk text
        t2 = time.time()
        chunks = chunk(text)
        t3 = time.time()

        if not chunks:
            return jsonify({'error': 'No text content found in PDF'}), 400

        avg_chunk_size = len(text) // len(chunks) if chunks else 0
        print(f"[CHUNK] {t3-t2:.2f}s, {len(chunks)} chunks (avg {avg_chunk_size} chars)")

        # Generate embeddings
        try:
            t4 = time.time()
            embeddings = [(c, embed(c)) for c in chunks]
            t5 = time.time()
            print(f"[EMBED] {t5-t4:.2f}s ({len(chunks)} embeddings)")
        except Exception as e:
            print(f"Error generating embeddings: {e}")
            return jsonify({'error': f'Failed to generate embeddings: {str(e)}'}), 500

        doc_id = str(uuid.uuid4())
        save_doc_embeddings(doc_id, embeddings)

        total_time = t5 - t0
        print(f"[UPLOAD DEBUG] Stored doc_id: {doc_id}")
        print(f"[UPLOAD DEBUG] Embeddings count: {len(embeddings)}")
        if embeddings:
            first_chunk_text = embeddings[0][0]
            print(f"[UPLOAD DEBUG] First chunk: {first_chunk_text[:100]}... (len={len(first_chunk_text)})")
        print(f"[DONE] Total processing: {total_time:.2f}s")

        # Store document in database with user ownership
        user_id = session.get('user_id')
        file_ext = Path(file.filename).suffix.lower()
        file_size = len(content)
        models.create_document(
            doc_id=doc_id,
            user_id=user_id,
            filename=file.filename,
            file_type=file_ext.replace('.', ''),  # e.g., 'pdf', 'docx'
            full_text=text,
            file_size=file_size
        )
        print(f"[DB] Document saved to database with owner: {user_id}")

        # Store in session
        session['doc_id'] = doc_id
        session['filename'] = file.filename

        # Generate document analysis (async-like, non-blocking)
        try:
            print("[ANALYSIS] Generating document analysis...")
            analysis = generate_document_analysis(text, file.filename)
            print(f"[ANALYSIS] Analysis generated: {len(analysis)} chars")
            # Store analysis in session so frontend can retrieve after upload
            session['doc_analysis'] = analysis
        except Exception as e:
            print(f"[ANALYSIS] Failed: {e}")
            # Don't fail the upload if analysis fails
            session['doc_analysis'] = None

        return jsonify({'doc_id': doc_id})

    except Exception as e:
        import traceback
        print(f"Unexpected error in upload: {e}")
        print(traceback.format_exc())
        return jsonify({'error': f'Internal server error: {str(e)}'}), 500

@app.route('/chat', methods=['POST'])
@login_required
def chat():
    """Chat with document"""
    try:
        if client is None:
            return jsonify({'error': 'OpenRouter client not initialized'}), 503

        data = request.get_json()
        doc_id = data.get("doc_id")
        messages = data.get("messages", [])

        # Debug logging
        print(f"[CHAT DEBUG] Received doc_id: {doc_id}")
        print(f"[CHAT DEBUG] DB keys: {list(DB.keys())}")
        print(f"[CHAT DEBUG] doc_id in DB: {doc_id in DB}")

        try:
            doc_chunks = get_doc_embeddings(doc_id)
        except KeyError:
            print(f"[CHAT ERROR] Document {doc_id} not found in database")
            return jsonify({'error': 'Document not found. Please upload a file first.'}), 404

        # Log document info
        print(f"[CHAT DEBUG] Document has {len(doc_chunks)} chunks")
        if doc_chunks:
            first_chunk_preview = doc_chunks[0][0][:100] if doc_chunks[0][0] else "Empty chunk"
            print(f"[CHAT DEBUG] First chunk preview: {first_chunk_preview}...")

        if not messages:
            return jsonify({'error': 'No messages provided'}), 400

        if not messages:
            return jsonify({'error': 'No messages provided'}), 400

        query = messages[-1]["content"]
        query_lower = query.lower().strip()

        # Save user message to chat history
        try:
            models.save_chat_message(doc_id, 'user', query)
        except Exception as e:
            print(f"[CHAT] Failed to save user message: {e}")

        # Small talk detection
        greeting_phrases = ['hi', 'hello', 'hey', 'good morning', 'good afternoon', 'good evening', 'greetings', 'howdy', "what's up", 'sup', 'yo']
        thanks_phrases = ['thanks', 'thank you','appreciate it', 'appreciate']
        farewell_phrases = ['bye', 'goodbye', 'see you', 'farewell', 'take care', 'later']
        how_are_you_phrases = ['how are you', "how's it going", 'how do you do']

        is_greeting = any(g in query_lower for g in greeting_phrases)
        is_thanks = any(t in query_lower for t in thanks_phrases)
        is_farewell = any(f in query_lower for f in farewell_phrases)
        is_how_are_you = any(h in query_lower for h in how_are_you_phrases)

        is_small_talk = is_greeting or is_thanks or is_farewell or is_how_are_you

        if is_small_talk:
            debug_type = []
            if is_greeting: debug_type.append("greeting")
            if is_thanks: debug_type.append("thanks")
            if is_farewell: debug_type.append("farewell")
            if is_how_are_you: debug_type.append("how-are-you")
            safe_query = safe_str(query, 40)
            print(f"[CHAT] Detected {', '.join(debug_type)}: '{safe_query}...' - small talk mode")
        else:
            safe_query = safe_str(query, 50)
            print(f"[QUERY] Document query: '{safe_query}...' - using context")

        # Get context (skip for small talk)
        if is_small_talk:
            context = ""
        else:
            try:
                q_emb = embed(query)
            except Exception as e:
                return jsonify({'error': f'Failed to embed query: {str(e)}'}), 500

            try:
                doc_chunks = get_doc_embeddings(doc_id)
                scored = [(c, cosine(q_emb, e)) for c, e in doc_chunks]
            except Exception as e:
                return jsonify({'error': f'Failed to score chunks: {str(e)}'}), 500

            top = sorted(scored, key=lambda x: x[1], reverse=True)[:TOP_K_RESULTS]
            context = "\n\n".join([c[0] for c in top])

            # Debug context
            print(f"[CHAT DEBUG] Retrieved {len(top)} top chunks")
            print(f"[CHAT DEBUG] Context length: {len(context)} chars")
            if context:
                print(f"[CHAT DEBUG] Context preview: {context[:200]}...")
            else:
                print(f"[CHAT DEBUG] Context is EMPTY - check document content and embedding quality")

            # If context is empty or too short, provide a fallback message to the AI
            if not context or len(context.strip()) < 10:
                context = "[NOTE: No relevant content could be found in the document for this query. The document may be scanned, contain only images without OCR, or the query is too unrelated to the content.]"
                print(f"[CHAT WARN] Using fallback context due to insufficient retrieval")

        # Build system prompt and final messages
        print(f"[CHAT DEBUG] is_small_talk: {is_small_talk}")

        if is_small_talk:
            system_content = """You are a friendly, warm AI assistant helping with a document.

BE HUMAN-LIKE:
- Use natural, casual language like talking to a friend
- Vary your responses - don't repeat the same phrases
- Use contractions (I'm, you're, that's, don't)
- Keep responses short (usually 1-2 sentences)
- Use emojis occasionally (😊 👋 💡 ✅)
- Be enthusiastic but not overbearing
- Reference previous conversation naturally if there is any

GREETING EXAMPLES (vary these):
- "Hey! 👋 I'm here and ready to help. What's on your mind?"
- "Hello! I'm excited to dive into this document with you. What would you like to explore?"
- "Hi there! 😊 I've got your document loaded up. Fire away with any questions!"

THANKS: "You're welcome! 😊 Happy to help!" or "Anytime! That's what I'm here for."

FAREWELLS: "Take care! Feel free to come back anytime. 👋" or "Goodbye! I'll be here when you need help."

HOW ARE YOU: "I'm doing great, thanks for asking! 😄 Ready to help. What's up?"

IMPORTANT: Keep it natural and varied."""
            final_messages = [{"role": "system", "content": system_content}] + messages
        else:
            system_content = """You are an expert AI tutor and document assistant with deep analytical skills.

CORE PRINCIPLES:
- You are a knowledgeable study companion who makes complex material accessible
- Be precise, thorough, and educational while maintaining a friendly tone
- Use examples generously to illustrate abstract concepts
- Build on previous questions in the conversation

RESPONSE STRUCTURE:
1. **Direct Answer**: Start with a clear, concise answer to the question
2. **Explanation**: Break down the concept step by step
3. **Example**: Provide at least one concrete real-world example
4. **Reference**: Quote or paraphrase specific parts from the document
5. **Check Understanding**: Ask a follow-up question or suggest related topics

TECHNICAL EXPLANATION GUIDELINES:
- Define technical terms before using them
- Use analogies: "Think of X as..." or "It's like when you..."
- Compare unfamiliar concepts to everyday experiences
- When discussing processes, describe them in logical order
- For numerical data, put it in context: "10% means roughly 1 out of 10"

CONTEXT USAGE (CRITICAL):
- The document excerpts below are your ONLY source of information
- Scan ALL provided context chunks thoroughly before answering
- If the answer spans multiple chunks, synthesize them coherently
- If the exact answer isn't there, state: "That specific information isn't covered in the document, but I can explain [related topic that IS present]"
- NEVER invent details, statistics, or quotes not present in the context

WHAT TO AVOID:
- "I don't have access to the document" → You DO have access via the context provided
- "Based on general knowledge" → Use ONLY the document
- Making up examples not derivable from the document
- Assuming information not explicitly stated

If there is NO relevant information in the context for the question:
- Clearly state: "I couldn't find information about that in the document."
- Suggest 2-3 related topics that ARE covered that might be helpful
- Ask if the user would like an explanation of those instead

Document context:
---
{context}
---"""
            final_system_content = system_content.format(context=context)
            print(f"[CHAT DEBUG] Final system content length: {len(final_system_content)} chars")
            # Log first 300 chars to see if context is included (use safe_str to avoid encoding errors)
            safe_preview = safe_str(final_system_content[:300], 300)
            print(f"[CHAT DEBUG] System content preview: {safe_preview}...")
            final_messages = [{"role": "system", "content": final_system_content}] + messages

        def generate():
            try:
                res = client.chat.completions.create(
                    model=OPENROUTER_MODEL,
                    stream=True,
                    messages=final_messages,
                    temperature=0.7,
                    max_tokens=500 if is_small_talk else 1000
                )

                full_response = []
                for chunk in res:
                    if chunk.choices[0].delta.content:
                        content = chunk.choices[0].delta.content
                        full_response.append(content)
                        yield content

                # Save assistant response to chat history
                if full_response:
                    full_response_text = ''.join(full_response)
                    try:
                        models.save_chat_message(doc_id, 'assistant', full_response_text)
                    except Exception as e:
                        print(f"[CHAT] Failed to save assistant message: {e}")

                # Log response for debugging (use safe_str to avoid Windows console encoding errors with emojis)
                if is_small_talk and full_response:
                    safe_response = safe_str(full_response_text, 100)
                    print(f"[AI] Response: '{safe_response}...'")

            except Exception as e:
                error_msg = str(e)
                if "401" in str(e) or "unauthorized" in error_msg.lower():
                    yield "Error: Authentication failed. Check API key."
                elif "insufficient_quota" in error_msg.lower() or "402" in str(e):
                    yield "Error: Quota exceeded. Add credits to your account."
                elif "rate limit" in error_msg.lower() or "429" in str(e):
                    yield "Error: Rate limit exceeded. Please wait."
                elif "model not found" in error_msg.lower() or "404" in str(e):
                    yield f"Error: Model '{OPENROUTER_MODEL}' not found or not accessible."
                else:
                    yield f"Error: {error_msg}"

        return Response(stream_with_context(generate()), mimetype='text/plain')

    except Exception as e:
        print(f"Unexpected error in chat: {e}")
        return jsonify({'error': f'Chat error: {str(e)}'}), 500

@app.route('/quiz', methods=['POST'])
@login_required
def generate_quiz():
    """Generate quiz questions based on document content"""
    try:
        if client is None:
            return jsonify({'error': 'OpenRouter client not initialized'}), 503

        data = request.get_json()
        doc_id = data.get("doc_id")
        topic = data.get("topic", "")  # Optional: specific topic to quiz on

        try:
            doc_chunks = get_doc_embeddings(doc_id)
        except KeyError:
            return jsonify({'error': 'Document not found. Please upload a file first.'}), 404

        # Get relevant context chunks
        if topic:
            # If topic specified, get relevant chunks using embedding similarity
            try:
                q_emb = embed(topic)
                scored = [(c, cosine(q_emb, e)) for c, e in doc_chunks]
                top = sorted(scored, key=lambda x: x[1], reverse=True)[:TOP_K_RESULTS]
                context_chunks = [c[0] for c in top]
            except Exception as e:
                print(f"Error finding topic chunks: {e}")
                context_chunks = [c[0] for c in doc_chunks[:5]]  # fallback to first 5
        else:
            # Use top chunks overall (first few, or could random sample)
            context_chunks = [c[0] for c in doc_chunks[:5]]

        context = "\n\n".join(context_chunks)

        quiz_prompt = f"""Based on the following document content, generate 3-5 quiz questions to test understanding.

The questions should:
- Mix multiple choice and short answer formats
- Cover key concepts from the content
- Include answers after each question (format: "Answer: ...")
- Vary difficulty (mix recall and application)
- Be clear and unambiguous

If a specific topic is provided, focus questions on that topic.

Document content:
---
{context}
---

Generate the quiz questions:"""

        try:
            response = client.chat.completions.create(
                model=OPENROUTER_MODEL,
                messages=[
                    {"role": "system", "content": "You are an expert educator creating quiz questions that test deep understanding."},
                    {"role": "user", "content": quiz_prompt}
                ],
                temperature=0.7,
                max_tokens=1000
            )
            quiz = response.choices[0].message.content
            return jsonify({'quiz': quiz})
        except Exception as e:
            print(f"Quiz generation API error: {e}")
            return jsonify({'error': f'Quiz generation failed: {str(e)}'}), 500

    except Exception as e:
        print(f"Quiz generation error: {e}")
        return jsonify({'error': f'Quiz error: {str(e)}'}), 500

@app.route('/document-analysis', methods=['POST'])
@login_required
def document_analysis():
    """Generate analysis for an uploaded document"""
    try:
        if client is None:
            return jsonify({'error': 'OpenRouter client not initialized'}), 503

        data = request.get_json()
        doc_id = data.get("doc_id")

        try:
            chunks = get_doc_embeddings(doc_id)
        except KeyError:
            return jsonify({'error': 'Document not found'}), 404

        # Reconstruct full text from chunks
        full_text = "\n\n".join([c[0] for c in chunks])

        # Get filename from session
        filename = session.get('filename', 'Unknown Document')

        # Generate analysis
        analysis = generate_document_analysis(full_text, filename)

        # Save analysis as first chat message (system-like) to history
        models.save_chat_message(doc_id, 'assistant', analysis)

        # Store in session so frontend can retrieve
        session['doc_analysis'] = analysis

        return jsonify({'analysis': analysis})

    except Exception as e:
        print(f"Document analysis error: {e}")
        return jsonify({'error': f'Analysis failed: {str(e)}'}), 500

# ==================== PROFILE & HISTORY ENDPOINTS ====================

@app.route('/profile', methods=['GET'])
@login_required
def get_profile():
    """Get current user's profile"""
    user_id = session.get('user_id')
    print(f"[PROFILE] Fetching profile for user_id: {user_id}")
    user = models.get_user_by_id(user_id)
    if not user:
        print(f"[PROFILE] User not found: {user_id}")
        return jsonify({'error': 'User not found'}), 404

    # Return sanitized user data
    print(f"[PROFILE] Returning: id={user['id']}, display_name={user['display_name']}, avatar_filename={user['avatar_filename']}")
    return jsonify({
        'id': user['id'],
        'email': user['email'],
        'display_name': user['display_name'],
        'avatar_url': models.get_avatar_url(user['id'], user['avatar_filename']) if user['avatar_filename'] else None,
        'created_at': user['created_at']
    })

@app.route('/profile/update', methods=['POST'])
@login_required
def update_profile():
    """Update user's display name"""
    user_id = session.get('user_id')
    if not user_id:
        print("[PROFILE] No user_id in session")
        return jsonify({'error': 'Not authenticated'}), 401

    data = request.get_json()
    if not data:
        print("[PROFILE] No JSON data in request")
        return jsonify({'error': 'Invalid request'}), 400

    display_name = data.get('display_name', '').strip()
    if not display_name:
        return jsonify({'error': 'Display name is required'}), 400

    print(f"[PROFILE] Updating name for user {user_id} to '{display_name}'")
    try:
        models.update_user_name(user_id, display_name)
        session['user_name'] = display_name  # Update session too
        print(f"[PROFILE] Successfully updated name for {user_id}")
        return jsonify({'success': True, 'message': 'Profile updated'})
    except Exception as e:
        print(f"[PROFILE] Error updating name for {user_id}: {e}")
        return jsonify({'error': f'Failed to update profile: {str(e)}'}), 500

@app.route('/profile/avatar', methods=['POST'])
@login_required
def upload_avatar():
    """Upload user avatar image"""
    user_id = session.get('user_id')

    if 'avatar' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['avatar']

    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    # Validate file type
    allowed_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.webp'}
    ext = Path(file.filename).suffix.lower()
    if ext not in allowed_extensions:
        return jsonify({'error': 'Only image files (PNG, JPG, GIF, WebP) are allowed'}), 400

    # Check file size (2MB max)
    file.seek(0, 2)  # Seek to end
    size = file.tell()
    file.seek(0)  # Reset
    if size > 2 * 1024 * 1024:
        return jsonify({'error': 'Avatar must be less than 2MB'}), 400

    try:
        # Delete old avatar if exists
        models.delete_avatar(user_id)

        # Save new avatar
        from PIL import Image
        import io

        # Open and resize to 200x200
        image = Image.open(io.BytesIO(file.read()))
        if image.mode != 'RGB':
            image = image.convert('RGB')
        image = image.resize((200, 200), Image.Resampling.LANCZOS)

        # Encode image to bytes in the correct format
        format_map = {
            '.png': 'PNG',
            '.jpg': 'JPEG',
            '.jpeg': 'JPEG',
            '.gif': 'GIF',
            '.webp': 'WEBP'
        }
        pil_format = format_map.get(ext, 'PNG')
        buffer = io.BytesIO()
        image.save(buffer, format=pil_format)
        image_bytes = buffer.getvalue()

        # Save to avatar directory
        avatar_filename = models.save_avatar(user_id, image_bytes, f"avatar{ext}")
        print(f"[AVATAR] Saved avatar file: {avatar_filename}")

        # Update user record
        models.update_user_avatar(user_id, avatar_filename)
        print(f"[AVATAR] Updated user {user_id} with avatar {avatar_filename}")

        return jsonify({
            'success': True,
            'avatar_url': models.get_avatar_url(user_id, avatar_filename),
            'message': 'Avatar uploaded successfully'
        })

    except Exception as e:
        print(f"Avatar upload error: {e}")
        return jsonify({'error': f'Failed to upload avatar: {str(e)}'}), 500

@app.route('/user/avatar/<user_id>')
def serve_avatar(user_id):
    """Serve user avatar image (protected)"""
    # Check if user is logged in and requesting their own avatar
    # For simplicity, we'll serve avatars publicly if you know the user_id
    # In production, add proper access control
    user = models.get_user_by_id(user_id)
    if not user or not user['avatar_filename']:
        # Return default avatar
        return '', 404

    avatar_filename = user['avatar_filename']

    # Check if it's a Blob URL
    if avatar_filename.startswith('http'):
        # Avatar stored in Vercel Blob
        try:
            import blob_storage
            blob_data = blob_storage.get_blob_file(user_id, avatar_filename)
            if blob_data:
                # Determine MIME type from URL or content-type
                ext = Path(avatar_filename).suffix.lower()
                mime_types = {
                    '.png': 'image/png',
                    '.jpg': 'image/jpeg',
                    '.jpeg': 'image/jpeg',
                    '.gif': 'image/gif',
                    '.webp': 'image/webp'
                }
                mime_type = mime_types.get(ext, 'image/jpeg')
                return Response(blob_data, mimetype=mime_type)
        except Exception as e:
            print(f"[AVATAR] Failed to fetch blob {avatar_filename}: {e}")
            return '', 404
        return '', 404
    else:
        # Local file storage
        from pathlib import Path
        avatar_path = models.get_avatar_path(user_id, avatar_filename)
        if not avatar_path.exists():
            return '', 404

        # Determine MIME type
        ext = avatar_path.suffix.lower()
        mime_types = {
            '.png': 'image/png',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.gif': 'image/gif',
            '.webp': 'image/webp'
        }
        mime_type = mime_types.get(ext, 'image/jpeg')

        return send_file(avatar_path, mimetype=mime_type)

@app.route('/history/documents', methods=['GET'])
@login_required
def get_document_history():
    """Get user's document history"""
    user_id = session.get('user_id')
    documents = models.get_user_documents(user_id)
    return jsonify({'documents': documents})

@app.route('/history/chat/<doc_id>', methods=['GET'])
@login_required
def get_chat_history(doc_id):
    """Get chat history for a specific document"""
    user_id = session.get('user_id')
    # Verify ownership
    doc = models.get_document(doc_id, user_id)
    if not doc:
        return jsonify({'error': 'Document not found or access denied'}), 404

    messages = models.get_chat_history(doc_id)
    return jsonify({'messages': messages})

@app.route('/documents/load', methods=['POST'])
@login_required
def load_document_from_history():
    """Load a document from history into the current session for continued Q&A"""
    user_id = session.get('user_id')
    data = request.get_json()
    doc_id = data.get("doc_id")

    # Get document from DB and verify ownership
    doc = models.get_document(doc_id, user_id)
    if not doc:
        return jsonify({'error': 'Document not found or access denied'}), 404

    # Reconstruct chunks from full_text
    full_text = doc['full_text']
    chunks = chunk(full_text)

    if not chunks:
        return jsonify({'error': 'Failed to reconstruct document chunks'}), 500

    # Generate embeddings
    try:
        embeddings = [(c, embed(c)) for c in chunks]
    except Exception as e:
        return jsonify({'error': f'Failed to generate embeddings: {str(e)}'}), 500

    # Store embeddings in database (and in memory cache)
    save_doc_embeddings(doc_id, embeddings)

    # Set session
    session['doc_id'] = doc_id
    session['filename'] = doc['filename']

    return jsonify({
        'success': True,
        'doc_id': doc_id,
        'filename': doc['filename'],
        'chunks_count': len(chunks)
    })

@app.route('/documents/delete', methods=['DELETE'])
@login_required
def delete_document():
    """Delete a document and its chat history"""
    user_id = session.get('user_id')
    doc_id = request.args.get('doc_id') or request.get_json().get('doc_id')

    if not doc_id:
        return jsonify({'error': 'Document ID is required'}), 400

    # Check if this is the currently loaded document
    if session.get('doc_id') == doc_id:
        # Clear from memory and session
        if doc_id in DB:
            del DB[doc_id]
        session.pop('doc_id', None)
        session.pop('filename', None)

    # Delete from database (document, chat history cascade, and embeddings)
    deleted = models.delete_document(doc_id, user_id)
    if not deleted:
        return jsonify({'error': 'Document not found or access denied'}), 404

    # Also delete embeddings (in case not cascade or from memory)
    delete_doc_embeddings(doc_id)

    return jsonify({'success': True, 'message': 'Document deleted'})

# ==================== STATIC FILE ROUTES ====================
# Flask automatically serves from /static and /templates folders
# Note: We import send_file for avatar serving
from flask import send_file

# Flask automatically serves from /static and /templates folders

# Initialize OpenRouter client at startup (for WSGI/vercel, not just __main__)
try:
    initialize_openrouter()
except Exception as e:
    print(f"[WARN] OpenRouter initialization failed: {e}")
    print("[WARN] Chat functionality will be disabled until OPENROUTER_API_KEY is properly set.")

# ==================== MAIN ======================

if __name__ == '__main__':
    initialize_openrouter()
    check_tesseract()  # Verify Tesseract installation
    print(f"\n[START] Unified Flask App")
    print(f"   Mode: {ENV_MODE}")
    print(f"   Debug: {debug_mode}")
    print(f"   Model: {OPENROUTER_MODEL}")
    print(f"\n   Open browser: http://localhost:5000")
    print(f"   Press Ctrl+C to stop\n")

    app.run(debug=debug_mode, port=5000, host='0.0.0.0')
